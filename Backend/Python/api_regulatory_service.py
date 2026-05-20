"""
Camada de serviço do fluxo regulatório guiado do BridgeMedAI.

Responsável por conduzir uma conversa estruturada em três passos:

    Step 1 — Análise do dispositivo:
        O utilizador descreve o dispositivo médico. O assistente devolve:
        - tipo e finalidade prevista;
        - classe de risco MDR (Anexo VIII) com justificação da regra;
        - enquadramento no AI Act e categoria de risco;
        - obrigações MDR aplicáveis;
        - normas harmonizadas relevantes;
        - distinção pré-mercado / pós-mercado;
        - pergunta final sobre preenchimento do documento PMCF.

    Step 2 — Recolha de informação em falta:
        Se o utilizador confirmar, o sistema carrega o template PMCF,
        identifica os campos ainda em falta (considerando o que já foi
        recolhido na conversa) e coloca perguntas agrupadas logicamente.

    Step 3 — Preenchimento e exportação:
        Com os dados recolhidos, o sistema preenche o template. Campos que
        exigem juízo humano, dados internos ou assinaturas são deixados com
        um aviso visível "⚠️ Preencher manualmente". No final é devolvido um
        caminho de download do ficheiro final (.docx).

Este módulo é deliberadamente independente do `api_rag_service` para não
afetar o fluxo RAG existente. A lógica HTTP vive em `api_main.py`.
"""

from __future__ import annotations

import json
import os
import re
import uuid
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Tuple

import ollama
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv


# ---------------------------------------------------------------------------
# Configuração e caminhos
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent.parent
ENV_PATH = PROJECT_ROOT / "Backend" / ".env"

load_dotenv(dotenv_path=ENV_PATH)

OLLAMA_REGULATORY_MODEL = os.getenv("OLLAMA_REGULATORY_MODEL", "qwen2.5:7b-instruct")

PMCF_TEMPLATE_PATH = (
    PROJECT_ROOT / os.getenv("PMCF_TEMPLATE_PATH", "Backend/templates/pmcf_template.docx")
).resolve()

PMCF_OUTPUT_DIR = (
    PROJECT_ROOT / os.getenv("PMCF_OUTPUT_DIR", "Backend/generated")
).resolve()
PMCF_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

MANUAL_FLAG_PREFIX = "⚠️ Preencher manualmente"


# ---------------------------------------------------------------------------
# Definição dos campos do template PMCF
# ---------------------------------------------------------------------------
# Cada campo estruturado é localizado numa célula (tabela, linha, coluna).
# O campo "auto" indica se o sistema deve tentar inferir o valor a partir da
# descrição do dispositivo e análise regulatória da Step 1. Os restantes são
# sempre perguntados ao utilizador.
#
# O campo "manual_only" força a flag de preenchimento manual (ex: assinaturas).

STRUCTURED_FIELDS: List[Dict[str, Any]] = [
    # --- Identificação do plano ---
    {"key": "pmcf_plan_number", "label": "Número do Plano PMCF",
     "table": 1, "row": 0, "col": 1, "group": "identificação", "auto": False},
    {"key": "pmcf_plan_date", "label": "Data do Plano PMCF (AAAA-MM-DD)",
     "table": 1, "row": 1, "col": 1, "group": "identificação", "auto": False},

    # --- Histórico de versões ---
    {"key": "version_number", "label": "Número da versão (ex: 1)",
     "table": 2, "row": 1, "col": 0, "group": "identificação", "auto": False},
    {"key": "version_date", "label": "Data de validade da versão (AAAA-MM-DD)",
     "table": 2, "row": 1, "col": 1, "group": "identificação", "auto": False},
    {"key": "version_created_by", "label": "Criado por (nome + assinatura digital)",
     "table": 2, "row": 1, "col": 2, "group": "identificação", "auto": False, "manual_only": True},
    {"key": "version_approved_by", "label": "Aprovado por (nome + assinatura digital)",
     "table": 2, "row": 1, "col": 3, "group": "identificação", "auto": False, "manual_only": True},
    {"key": "version_description", "label": "Descrição da alteração",
     "table": 2, "row": 1, "col": 4, "group": "identificação", "auto": False},

    # --- Fabricante ---
    {"key": "manufacturer_name", "label": "Nome do fabricante",
     "table": 3, "row": 0, "col": 1, "group": "fabricante", "auto": False},
    {"key": "manufacturer_address", "label": "Morada do fabricante",
     "table": 3, "row": 1, "col": 1, "group": "fabricante", "auto": False},
    {"key": "manufacturer_srn", "label": "Single Registration Number (SRN)",
     "table": 3, "row": 2, "col": 1, "group": "fabricante", "auto": False},
    {"key": "manufacturer_prrc", "label": "Pessoa Responsável pela Conformidade Regulamentar (PRRC) — Artigo 15.º MDR",
     "table": 3, "row": 3, "col": 1, "group": "fabricante", "auto": False},
    {"key": "manufacturer_email", "label": "Email de contacto",
     "table": 3, "row": 4, "col": 1, "group": "fabricante", "auto": False},
    {"key": "manufacturer_phone", "label": "Telefone de contacto",
     "table": 3, "row": 5, "col": 1, "group": "fabricante", "auto": False},

    # --- Produto ---
    {"key": "product_name", "label": "Nome comercial do produto",
     "table": 4, "row": 1, "col": 0, "group": "produto", "auto": True},
    {"key": "product_version", "label": "Versão do produto",
     "table": 4, "row": 1, "col": 1, "group": "produto", "auto": False},
    {"key": "product_sw_version", "label": "Versão de software",
     "table": 4, "row": 1, "col": 2, "group": "produto", "auto": False},
    {"key": "product_udi", "label": "Basic UDI-DI (se disponível)",
     "table": 4, "row": 1, "col": 3, "group": "produto", "auto": False},
    {"key": "product_cert_number", "label": "Número de Certificado CE (se aplicável)",
     "table": 4, "row": 1, "col": 4, "group": "produto", "auto": False},
    {"key": "product_cndn", "label": "Código CNDN (ou equivalente)",
     "table": 4, "row": 1, "col": 5, "group": "produto", "auto": False},
    {"key": "product_class_rule", "label": "Classe e Regra MDR aplicável",
     "table": 4, "row": 1, "col": 6, "group": "produto", "auto": True},
]


# Secções de texto livre. Cada uma é localizada pela pesquisa do seu heading
# no documento; o conteúdo preenchido é inserido numa linha nova imediatamente
# após o heading, assinalado como "Conteúdo preenchido".
#
# Os campos "auto": True são os que a Step 1 tenta inferir a partir da
# descrição do dispositivo.
NARRATIVE_SECTIONS: List[Dict[str, Any]] = [
    {"key": "intended_purpose",     "heading": "Intended Purpose",
     "label": "Finalidade prevista (copiar palavra-a-palavra da documentação técnica)", "auto": True},
    {"key": "intended_users",       "heading": "Intended Users",
     "label": "Utilizadores pretendidos", "auto": True},
    {"key": "patient_population",   "heading": "Patient Population",
     "label": "População de doentes alvo", "auto": True},
    {"key": "medical_indication",   "heading": "Medical Indication",
     "label": "Indicações clínicas", "auto": True},
    {"key": "contraindications",    "heading": "Contraindications",
     "label": "Contraindicações", "auto": True},
    {"key": "variants",             "heading": "List and description of any variants and/or configurations covered in PMF",
     "label": "Variantes e configurações cobertas (se não aplicável, indicar)", "auto": False},
    {"key": "accessories",          "heading": "List of accessories covered in PMCF",
     "label": "Acessórios cobertos (se não aplicável, indicar)", "auto": False},
    {"key": "expected_lifetime",    "heading": "Expected Lifetime",
     "label": "Vida útil esperada", "auto": True},
    {"key": "novel_product",        "heading": "Novel product",
     "label": "O produto é novo no mercado? Justificar", "auto": True},
    {"key": "novel_procedure",      "heading": "Novel related clinical procedure",
     "label": "Corresponde a um procedimento clínico novo? Justificar", "auto": True},
    {"key": "novel_features",       "heading": "Novel features",
     "label": "Existem características novas? Descrever", "auto": True},
    {"key": "pmcf_objectives",      "heading": "Objectives",
     "label": "Objetivos do PMCF (por cada risco/benefício, uma hipótese mensurável)", "auto": True},
    {"key": "pmcf_methods",         "heading": "Methods and Procedures",
     "label": "Métodos e procedimentos PMCF (rever literatura, inquéritos, registos, testes internos, estudo PMCF)", "auto": True},
    {"key": "estimated_report_date","heading": "Estimate date of the PMCF Report",
     "label": "Data estimada do relatório PMCF (AAAA-MM-DD; no máximo anual)", "auto": False},
]


def all_field_keys() -> List[str]:
    return [f["key"] for f in STRUCTURED_FIELDS] + [s["key"] for s in NARRATIVE_SECTIONS]


def field_by_key(key: str) -> Optional[Dict[str, Any]]:
    for f in STRUCTURED_FIELDS:
        if f["key"] == key:
            return f
    for s in NARRATIVE_SECTIONS:
        if s["key"] == key:
            return s
    return None


# ---------------------------------------------------------------------------
# Estado de sessão
# ---------------------------------------------------------------------------
@dataclass
class RegulatorySession:
    session_id: str
    step: str = "awaiting_description"  # awaiting_description | awaiting_fill_confirmation | collecting_info | document_ready
    device_description: str = ""
    analysis_text: str = ""          # resposta narrativa em markdown da Step 1
    analysis_structured: Dict[str, Any] = field(default_factory=dict)  # metadados estruturados
    collected: Dict[str, str] = field(default_factory=dict)  # {field_key: value}
    missing_keys: List[str] = field(default_factory=list)    # pendentes na Step 2
    asked_keys: List[str] = field(default_factory=list)      # já perguntadas
    template_path: str = str(PMCF_TEMPLATE_PATH)
    generated_path: Optional[str] = None
    flagged_fields: List[Dict[str, str]] = field(default_factory=list)
    history: List[Dict[str, str]] = field(default_factory=list)  # [{role, content}]


_SESSIONS: Dict[str, RegulatorySession] = {}
_SESSIONS_LOCK = Lock()


def get_or_create_session(session_id: Optional[str]) -> RegulatorySession:
    with _SESSIONS_LOCK:
        if session_id and session_id in _SESSIONS:
            return _SESSIONS[session_id]
        sid = session_id or str(uuid.uuid4())
        session = RegulatorySession(session_id=sid)
        _SESSIONS[sid] = session
        return session


def get_session(session_id: str) -> Optional[RegulatorySession]:
    with _SESSIONS_LOCK:
        return _SESSIONS.get(session_id)


# ---------------------------------------------------------------------------
# Prompts base
# ---------------------------------------------------------------------------
SYSTEM_PROMPT_STEP1 = """És um assistente regulatório especialista em dispositivos médicos da UE, do projeto BridgeMedAI.

REGRAS OBRIGATÓRIAS:
- Responde sempre em português de Portugal, num registo técnico e conciso.
- Nunca inventes artigos, anexos, regras, normas ou considerandos. Cita apenas referências reais do MDR (Regulamento (UE) 2017/745), do AI Act (Regulamento (UE) 2024/1689), ISO, IEC e MDCG/MEDDEV.
- Nunca saltes a etapa de classificação de risco. Usa sempre o MDR Anexo VIII.
- Quando não tiveres informação suficiente, indica-o explicitamente em vez de inventar.
- Termina sempre a resposta perguntando ao utilizador se pretende que preenchas o Plano de Follow-Up Clínico Pós-Mercado (PMCF).


REGRAS DE CONTEÚDO REGULATÓRIO:
- Nunca digas que a vigilância pós-comercialização/PMS não é obrigatória para um dispositivo médico MDR. Deve ser apresentada como obrigação pós-mercado proporcional ao tipo, classe e risco do dispositivo.
- Nunca digas que o PMCF é “não aplicável” apenas por causa da classe do dispositivo. O PMCF faz parte da avaliação clínica e do sistema PMS; pode ser justificado que não sejam necessárias atividades PMCF específicas, mas essa justificação depende da avaliação clínica, risco, novidade, dados clínicos existentes e desempenho pós-mercado.
- Para PMCF, usa formulações como: “PMCF/ACPC deve ser planeado ou justificado no âmbito da avaliação clínica e PMS; a extensão das atividades depende do risco, novidade e evidência clínica disponível.”
- Para QMS, não digas simplesmente “ISO 13485 não é obrigatório”. Diz: “o fabricante deve ter um sistema de gestão da qualidade conforme o MDR; ISO 13485 pode ser usada como referência reconhecida para estruturar esse sistema.”
- Para gestão de risco, apresenta ISO 14971 como norma prática/referência técnica, sem inventar obrigatoriedade legal direta se o contexto não a sustentar.
- Para AI Act, se o dispositivo descrito não tiver IA, diz: “não foi descrito um sistema de IA; por isso, não há enquadramento AI Act suficiente nesta descrição.”
- Para AI Act, se houver IA integrada num dispositivo médico, separa sempre: classe MDR ≠ categoria de risco AI Act.
- Não cites artigos do AI Act se não estiveres seguro do artigo aplicável.
- Não uses expressões absolutas como “não abrangido pelo AI Act” quando apenas falta descrição de IA. Prefere: “não há indícios suficientes de sistema de IA na descrição fornecida.”


REGRAS CANÓNICAS DE CLASSIFICAÇÃO MDR:
- Um termómetro digital com finalidade médica de medição da temperatura corporal deve ser tratado como dispositivo ativo de diagnóstico/monitorização; classe provável: Classe IIa pela Regra 10 do Anexo VIII, salvo se a descrição indicar apenas uso não médico/bem-estar.
- Não classifiques um termómetro digital médico como Classe I apenas por ser não invasivo. Sendo ativo e usado para diagnóstico/monitorização, considera primeiro a Regra 10.
- Um termómetro sem finalidade médica prevista pode ficar fora do MDR, mas isso depende dos claims/finalidade prevista do fabricante.
- Se houver algoritmo de IA que estima temperatura central, explica separadamente: classe MDR provável pela Regra 10 e eventual análise AI Act apenas se a funcionalidade for realmente um sistema de IA.

FORMATO DE RESPOSTA:
Devolves SEMPRE duas secções separadas pela linha exata ---JSON---

A primeira é a análise em texto simples com exatamente estas subsecções numeradas. NÃO uses asteriscos duplos (**texto**), NÃO uses cardinais (# ## ###) nem sublinhados para formatação. Começa cada secção por "N. Título — descrição" (com hífen longo) e, dentro de cada secção, usa parágrafos ou listas com "- " (hífen). Estrutura obrigatória:
1. Tipo de dispositivo e finalidade prevista
2. Classificação MDR (Anexo VIII) — indica Classe e Regra(s) aplicável(is) com justificação
3. Enquadramento no AI Act — indica se é abrangido e a categoria de risco (proibido / alto risco / risco limitado / risco mínimo), citando o(s) artigo(s) relevante(s)
4. Obrigações MDR aplicáveis — avaliação da conformidade, documentação técnica, QMS (ISO 13485), gestão de risco (ISO 14971), vigilância pós-mercado (PMS), plano de PMCF, UDI, registo em EUDAMED, etc.
5. Normas e orientações aplicáveis — ISO 13485, ISO 14971, IEC 62304, IEC 62366-1, ISO 14155, MDCG 2019-11, MEDDEV 2.7/1 rev.4, MEDDEV 2.12/2 rev.2, etc., conforme relevantes
6. Pré-mercado vs pós-mercado — separa claramente as obrigações antes da colocação no mercado das aplicáveis após a colocação
7. Pergunta final ao utilizador — escreve apenas a pergunta: Queres que preencha o Plano de Follow-Up Clínico Pós-Mercado (PMCF) com base nesta análise?

A segunda secção é um objeto JSON válido (sem markdown à volta) com a seguinte estrutura exata:
{
  "device_type": "tipo genérico do dispositivo",
  "intended_purpose": "finalidade prevista, uma ou duas frases",
  "mdr_class": "I | IIa | IIb | III",
  "mdr_rule": "Regra X do Anexo VIII",
  "ai_act_applies": true|false,
  "ai_act_risk": "proibido|alto_risco|risco_limitado|risco_minimo|nao_aplicavel",
  "is_software": true|false,
  "intended_users": "descrição dos utilizadores pretendidos",
  "patient_population": "descrição da população de doentes",
  "medical_indication": "indicações clínicas principais",
  "contraindications": "contraindicações conhecidas ou 'A confirmar pelo fabricante'",
  "product_name": "nome comercial inferido ou 'A confirmar'",
  "product_class_rule": "ex: Classe IIa (Regra 11 do Anexo VIII do MDR)",
  "expected_lifetime": "vida útil estimada ou 'A confirmar pelo fabricante'",
  "novel_product": "é produto novo? sim/não + justificação",
  "novel_procedure": "envolve procedimento clínico novo? sim/não + justificação",
  "novel_features": "características novas relevantes ou 'Nenhuma identificada'",
  "pmcf_objectives": "objetivos PMCF mensuráveis propostos com base nos riscos/benefícios",
  "pmcf_methods": "métodos PMCF propostos (revisão literatura, inquéritos, registo, testes internos, estudo PMCF)"
}

Usa valores curtos mas concretos. Se não houver informação suficiente num campo, escreve "A confirmar pelo fabricante"."""


SYSTEM_PROMPT_STEP2_MAP = """És um assistente regulatório a interpretar as respostas do fabricante para preencher um Plano PMCF.

Recebes:
- a lista de perguntas que foram feitas (com chaves internas);
- a resposta em texto livre do fabricante.

Tens de devolver APENAS um objeto JSON válido (sem markdown, sem comentários) no qual, para cada chave pedida, atribuis o valor correspondente extraído da resposta. Regras:
- Se o utilizador deixou uma pergunta sem resposta clara, usa a string exata "NAO_RESPONDIDO".
- Se o utilizador indicou que o campo não se aplica, usa "NAO_APLICAVEL".
- Não inventes valores. Não acrescentes comentários nem explicações fora do JSON.
- Em datas, devolve formato AAAA-MM-DD se possível.
- Em português de Portugal."""


# ---------------------------------------------------------------------------
# Utilitários de LLM
# ---------------------------------------------------------------------------
def _chat(messages: List[Dict[str, str]], temperature: float = 0.2) -> str:
    try:
        response = ollama.chat(
            model=OLLAMA_REGULATORY_MODEL,
            messages=messages,
            options={
                "temperature": temperature,
                "num_ctx": 4096,
                "num_predict": 1200,
            },
            keep_alive="10m",
            stream=False,
        )
        return response["message"]["content"]
    except Exception as exc:
        raise RuntimeError(
            f"Falha ao executar o modelo regulatório '{OLLAMA_REGULATORY_MODEL}': {exc}"
        ) from exc


def _extract_json(text: str) -> Dict[str, Any]:
    """Extrai o primeiro objeto JSON plausível de uma resposta do modelo."""
    if not text:
        return {}
    # procurar blocos ```json ... ```
    fenced = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL | re.IGNORECASE)
    if fenced:
        candidate = fenced.group(1)
    else:
        # fallback: maior objeto {...}
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return {}
        candidate = text[start:end + 1]
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        # última tentativa: remover vírgulas finais
        cleaned = re.sub(r",\s*([}\]])", r"\1", candidate)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            return {}


# ---------------------------------------------------------------------------
# Step 1 — análise
# ---------------------------------------------------------------------------
def analyze_device(session_id: Optional[str], description: str) -> Dict[str, Any]:
    """Executa a Step 1 e prepara a sessão para a Step 2."""
    description = (description or "").strip()
    if not description:
        raise ValueError("A descrição do dispositivo não pode estar vazia.")

    session = get_or_create_session(session_id)
    session.device_description = description
    session.history.append({"role": "user", "content": description})

    raw = _chat(
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT_STEP1},
            {"role": "user", "content": f"Descrição do dispositivo do fabricante:\n\n{description}"},
        ],
        temperature=0.2,
    )

    narrative, structured = _split_step1_response(raw)

    session.analysis_text = narrative
    session.analysis_structured = structured
    session.step = "awaiting_fill_confirmation"

    # pré-preenche os campos "auto" a partir da análise estruturada
    _seed_collected_from_analysis(session)

    session.history.append({"role": "assistant", "content": narrative})

    return {
        "session_id": session.session_id,
        "step": session.step,
        "assistant_text": narrative,
        "analysis": structured,
        "pending_action": "confirm_fill_pmcf",
    }


def _split_step1_response(raw: str) -> Tuple[str, Dict[str, Any]]:
    """Separa a resposta do modelo entre o texto em markdown e o JSON final."""
    if "---JSON---" in raw:
        narrative, _, json_part = raw.partition("---JSON---")
        structured = _extract_json(json_part)
    else:
        # o modelo pode ter devolvido apenas JSON ou apenas texto
        structured = _extract_json(raw)
        if structured:
            # remove o bloco JSON do narrativo
            narrative = re.sub(r"```(?:json)?\s*\{.*?\}\s*```", "", raw, flags=re.DOTALL)
            narrative = re.sub(r"\{[\s\S]*\}\s*$", "", narrative)
        else:
            narrative = raw
    return narrative.strip(), structured or {}


def _seed_collected_from_analysis(session: RegulatorySession) -> None:
    """Copia os campos derivados automaticamente da análise para os recolhidos."""
    s = session.analysis_structured
    desc = (session.device_description or "").lower()

    product_class_rule = s.get("product_class_rule")

    if not product_class_rule:
        mdr_class = s.get("mdr_class")
        mdr_rule = s.get("mdr_rule")
        if mdr_class and mdr_rule:
            product_class_rule = f"Classe {mdr_class} ({mdr_rule})"

    # fallback específico para termómetro ativo de monitorização clínica
    if not product_class_rule:
        if (
            ("termómetro" in desc or "termometro" in desc)
            and ("ativo" in desc or "digital" in desc)
            and (
                "monitorização" in desc
                or "monitorizacao" in desc
                or "sinais vitais" in desc
                or "temperatura corporal" in desc
            )
        ):
            product_class_rule = "Classe IIa provável (Regra 10 do Anexo VIII do MDR) — a confirmar conforme finalidade prevista e criticidade da monitorização."

    mapping = {
        "intended_purpose": s.get("intended_purpose"),
        "intended_users": s.get("intended_users"),
        "patient_population": s.get("patient_population"),
        "medical_indication": s.get("medical_indication"),
        "contraindications": s.get("contraindications"),
        "product_name": s.get("product_name"),
        "product_class_rule": product_class_rule,
        "expected_lifetime": s.get("expected_lifetime"),
        "novel_product": s.get("novel_product"),
        "novel_procedure": s.get("novel_procedure"),
        "novel_features": s.get("novel_features"),
        "pmcf_objectives": s.get("pmcf_objectives"),
        "pmcf_methods": s.get("pmcf_methods"),
    }

    for key, val in mapping.items():
        if val and str(val).strip() and str(val).strip().lower() not in {
            "a confirmar pelo fabricante",
            "a confirmar",
            "não determinado",
            "nao determinado",
            "não aplicável",
            "nao aplicavel",
        }:
            session.collected[key] = str(val).strip()


# ---------------------------------------------------------------------------
# Step 2 — recolha de informação
# ---------------------------------------------------------------------------
GROUP_ORDER = ["identificação", "fabricante", "produto", "secções narrativas"]
GROUP_LABELS = {
    "identificação": "Identificação e versão do plano",
    "fabricante": "Fabricante",
    "produto": "Produto",
    "secções narrativas": "Conteúdo clínico / pós-mercado",
}


def _missing_fields_groups(session: RegulatorySession) -> Dict[str, List[Dict[str, str]]]:
    """Agrupa os campos que ainda não têm valor recolhido."""
    groups: Dict[str, List[Dict[str, str]]] = {g: [] for g in GROUP_ORDER}

    for f in STRUCTURED_FIELDS:
        key = f["key"]
        if key in session.collected:
            continue
        # campos manual_only ficam para Step 3, não se perguntam
        if f.get("manual_only"):
            continue
        groups[f["group"]].append({"key": key, "label": f["label"]})

    for s in NARRATIVE_SECTIONS:
        key = s["key"]
        if key in session.collected:
            continue
        groups["secções narrativas"].append({"key": key, "label": s["label"]})

    return {g: items for g, items in groups.items() if items}


def start_collection(session_id: str) -> Dict[str, Any]:
    """Chamado quando o utilizador aceita preencher o PMCF. Gera a lista inicial de perguntas."""
    session = get_session(session_id)
    if not session:
        raise ValueError("Sessão regulatória não encontrada.")
    if session.step not in {"awaiting_fill_confirmation", "collecting_info"}:
        raise ValueError(f"Passo inválido para iniciar recolha: {session.step}.")

    session.step = "collecting_info"
    groups = _missing_fields_groups(session)

    if not groups:
        # nada em falta — gerar directamente
        return generate_document(session_id)

    flat_keys: List[str] = []
    for g in GROUP_ORDER:
        for item in groups.get(g, []):
            flat_keys.append(item["key"])
    session.missing_keys = flat_keys
    session.asked_keys = list(flat_keys)

    message = _format_collection_message(session, groups)
    session.history.append({"role": "assistant", "content": message})

    return {
        "session_id": session.session_id,
        "step": session.step,
        "assistant_text": message,
        "missing_groups": groups,
        "pending_action": "answer_missing_info",
    }


def _format_collection_message(session: RegulatorySession, groups: Dict[str, List[Dict[str, str]]]) -> str:
    auto_count = sum(1 for k in session.collected)
    lines = [
        "Perfeito — vou preencher o Plano PMCF a partir do template carregado.",
        "",
        f"Já consegui extrair **{auto_count}** campos da tua descrição inicial (finalidade, utilizadores, indicações, classe, etc.).",
        "Para fechar o documento, preciso que respondas às perguntas abaixo. Podes responder tudo numa mensagem só — basta identificares a que campo responde cada resposta, ou podes seguir a ordem das perguntas.",
        "",
        "Campos que não souberes: escreve *não sei*, *não aplicável* ou deixa em branco — esses ficam marcados no documento como *\"⚠️ Preencher manualmente\"*.",
        "",
    ]
    for g in GROUP_ORDER:
        items = groups.get(g, [])
        if not items:
            continue
        lines.append(f"### {GROUP_LABELS[g]}")
        for item in items:
            lines.append(f"- {item['label']}")
        lines.append("")
    return "\n".join(lines).strip()


def collect_answers(session_id: str, user_text: str) -> Dict[str, Any]:
    """Recebe o texto de resposta e mapeia para as chaves em aberto via LLM."""
    session = get_session(session_id)
    if not session:
        raise ValueError("Sessão regulatória não encontrada.")
    if session.step != "collecting_info":
        raise ValueError(f"Passo inválido para recolher respostas: {session.step}.")

    user_text = (user_text or "").strip()
    if not user_text:
        raise ValueError("A resposta não pode estar vazia.")

    session.history.append({"role": "user", "content": user_text})

    pending_keys = [k for k in session.missing_keys if k not in session.collected]
    if not pending_keys:
        # nada pendente, gerar doc
        return generate_document(session_id)

    pending_fields = [
        {"key": k, "label": (field_by_key(k) or {}).get("label", k)}
        for k in pending_keys
    ]

    user_prompt = (
        "Perguntas ainda em aberto (chave → descrição):\n"
        + "\n".join([f"- {p['key']}: {p['label']}" for p in pending_fields])
        + "\n\nResposta do fabricante:\n\"\"\"\n"
        + user_text
        + "\n\"\"\"\n\nDevolve APENAS um objeto JSON em que cada chave pedida recebe o valor em texto."
    )

    raw = _chat(
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT_STEP2_MAP},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
    )
    mapped = _extract_json(raw)

    newly_set: List[str] = []
    skipped: List[str] = []

    for key in pending_keys:
        if key not in mapped:
            continue
        val = str(mapped[key]).strip()
        if not val:
            continue
        if val.upper() in {"NAO_RESPONDIDO", "NÃO_RESPONDIDO", "NAO RESPONDIDO"}:
            skipped.append(key)
            continue
        if val.upper() in {"NAO_APLICAVEL", "NÃO_APLICÁVEL", "NAO APLICAVEL"}:
            session.collected[key] = "Não aplicável."
            newly_set.append(key)
            continue
        session.collected[key] = val
        newly_set.append(key)

    still_missing = [k for k in pending_keys if k not in session.collected and k not in skipped]
    session.missing_keys = still_missing

    if still_missing:
        assistant_msg = _format_still_missing(newly_set, still_missing)
        session.history.append({"role": "assistant", "content": assistant_msg})
        return {
            "session_id": session.session_id,
            "step": session.step,
            "assistant_text": assistant_msg,
            "collected_count": len(session.collected),
            "still_missing": still_missing,
            "pending_action": "answer_missing_info",
        }

    # tudo que o utilizador queria responder foi respondido → gerar doc
    return generate_document(session_id)


def _format_still_missing(newly_set: List[str], still_missing: List[str]) -> str:
    lines = []
    if newly_set:
        lines.append(f"Registei **{len(newly_set)}** campos. Faltam ainda:")
    else:
        lines.append("Ainda não consegui mapear todas as respostas. Faltam:")
    for key in still_missing:
        f = field_by_key(key) or {}
        lines.append(f"- {f.get('label', key)}")
    lines.append("")
    lines.append("Podes responder a todos, a alguns, ou dizer *não sei* — nesses casos marco-os como *\"⚠️ Preencher manualmente\"* e geramos o documento.")
    return "\n".join(lines)


def skip_remaining_and_generate(session_id: str) -> Dict[str, Any]:
    """Salta as respostas em falta e gera o documento com as flags apropriadas."""
    session = get_session(session_id)
    if not session:
        raise ValueError("Sessão regulatória não encontrada.")
    session.missing_keys = []
    return generate_document(session_id)


# ---------------------------------------------------------------------------
# Step 3 — geração do documento
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _delete_table(table) -> None:
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _table_text(table) -> str:
    parts = []
    for row in table.rows:
        for cell in row.cells:
            parts.append(cell.text or "")
    return "\n".join(parts)


def _clear_cell_and_write(cell, text: str, flag: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if flag:
        run.font.color.rgb = RGBColor(0xA2, 0x2D, 0x2D)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _cleanup_generated_pmcf_doc(doc: Document) -> None:
    """
    Remove instruções, exemplos e placeholders do template antes de guardar
    o documento final.
    """

    # 1) Remover parágrafos de instrução do template
    delete_until_pmcf_identification = False

    for paragraph in list(doc.paragraphs):
        text = (paragraph.text or "").strip()
        text_lower = text.lower()

        if text == "TEMPLATE INSTRUCTIONS":
            delete_until_pmcf_identification = True
            _delete_paragraph(paragraph)
            continue

        if delete_until_pmcf_identification:
            if re.match(r"^1\.\s*PMCF Identification", text, flags=re.IGNORECASE):
                delete_until_pmcf_identification = False
            else:
                _delete_paragraph(paragraph)
            continue

        instruction_patterns = [
            "important:",
            "examples:",
            "for example:",
            "please note",
            "this template document contains guidance notes",
            "sections that are indicated as optional",
            "clean up the table as needed",
            "guidance documents that may further help",
            "see section",
            "see user profile",
            "see tmp-",
            "as per mdcg",
        ]

        if any(p in text_lower for p in instruction_patterns):
            _delete_paragraph(paragraph)
            continue

        # remover linhas só com placeholders/exemplos
        if re.fullmatch(r"(?i)\s*(xx|xxx|yyy|\[.*?\]|e\.g\..*)\s*", text):
            _delete_paragraph(paragraph)
            continue

    # 2) Remover tabelas de histórico do template e tabelas claramente exemplificativas
    for table in list(doc.tables):
        t = _table_text(table).lower()

        if (
            "template version history" in t
            or "first version" in t
            or "célia cruz" in t
            or "celia cruz" in t
        ):
            _delete_table(table)
            continue

    # 3) Substituir células ainda exemplificativas por aviso manual
    placeholder_patterns = [
        r"\bXX\b",
        r"\bXXX\b",
        r"\bYYY\b",
        r"\be\.g\.",
        r"\[Title of the test report\]",
        r"Statistical sample size estimation",
        r"Start of the survey:",
        r"End of the survey:",
        r"Data analysis:",
    ]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if not text.strip():
                    continue

                if MANUAL_FLAG_PREFIX in text or "Conteúdo preenchido:" in text:
                    continue

                if any(re.search(p, text, flags=re.IGNORECASE) for p in placeholder_patterns):
                    _clear_cell_and_write(
                        cell,
                        f"{MANUAL_FLAG_PREFIX} — rever e adaptar ao dispositivo",
                        flag=True,
                    )


def generate_document(session_id: str) -> Dict[str, Any]:
    """Preenche o template e devolve o caminho do ficheiro final."""
    session = get_session(session_id)
    if not session:
        raise ValueError("Sessão regulatória não encontrada.")

    template_path = Path(session.template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Template PMCF não encontrado em {template_path}.")

    doc = Document(str(template_path))

    filled_fields: List[Dict[str, str]] = []
    flagged_fields: List[Dict[str, str]] = []

    # --- Preencher células estruturadas ---
    for f in STRUCTURED_FIELDS:
        key = f["key"]
        value = session.collected.get(key)
        is_manual = bool(f.get("manual_only"))

        if is_manual or not value:
            reason = "campo requer assinatura/decisão humana" if is_manual else "informação não fornecida pelo utilizador"
            display = f"{MANUAL_FLAG_PREFIX} — {reason}"
            _set_table_cell(doc, f["table"], f["row"], f["col"], display, flag=True)
            flagged_fields.append({"key": key, "label": f["label"], "reason": reason})
        else:
            _set_table_cell(doc, f["table"], f["row"], f["col"], value, flag=False)
            filled_fields.append({"key": key, "label": f["label"]})

    # --- Preencher secções narrativas ---
    for s in NARRATIVE_SECTIONS:
        key = s["key"]
        value = session.collected.get(key)
        if value:
            _insert_after_heading(doc, s["heading"], f"Conteúdo preenchido: {value}", flag=False)
            filled_fields.append({"key": key, "label": s["label"]})
        else:
            _insert_after_heading(
                doc,
                s["heading"],
                f"{MANUAL_FLAG_PREFIX} — informação não fornecida pelo utilizador",
                flag=True,
            )
            flagged_fields.append({
                "key": key,
                "label": s["label"],
                "reason": "informação não fornecida pelo utilizador",
            })

    # --- Limpar instruções e exemplos do template ---
    _cleanup_generated_pmcf_doc(doc)

    # --- Guardar ficheiro ---
    out_name = f"PMCF_{session.session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = PMCF_OUTPUT_DIR / out_name
    doc.save(str(out_path))

    session.generated_path = str(out_path)
    session.flagged_fields = flagged_fields
    session.step = "document_ready"

    summary_text = _format_generation_summary(filled_fields, flagged_fields, out_name)
    session.history.append({"role": "assistant", "content": summary_text})

    return {
        "session_id": session.session_id,
        "step": session.step,
        "assistant_text": summary_text,
        "filled_fields": filled_fields,
        "flagged_fields": flagged_fields,
        "download_name": out_name,
        "download_url": f"/regulatory/download/{session.session_id}",
        "pending_action": "download_document",
    }


def _format_generation_summary(filled: List[Dict[str, str]], flagged: List[Dict[str, str]], filename: str) -> str:
    lines = [
        "O documento **Plano PMCF** foi preenchido e está pronto para descarregar.",
        "",
        f"- Campos preenchidos automaticamente: **{len(filled)}**",
        f"- Campos marcados para revisão manual: **{len(flagged)}**",
        "",
        "### Campos preenchidos",
    ]
    if filled:
        for item in filled:
            lines.append(f"- {item['label']}")
    else:
        lines.append("- (nenhum)")

    lines.append("")
    lines.append("### Campos a rever manualmente")
    if flagged:
        for item in flagged:
            lines.append(f"- ⚠️ **{item['label']}** — {item['reason']}")
    else:
        lines.append("- (nenhum)")

    lines.append("")
    lines.append(f"📎 Ficheiro: `{filename}`")
    lines.append("")
    lines.append("Clica no botão *Descarregar PMCF* para guardar o ficheiro. Todos os campos marcados com ⚠️ precisam de ser revistos e preenchidos manualmente antes da assinatura final.")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Edição do DOCX
# ---------------------------------------------------------------------------
def _set_table_cell(doc: Document, table_idx: int, row_idx: int, col_idx: int, text: str, flag: bool) -> None:
    try:
        cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
    except (IndexError, AttributeError):
        return
    # limpar conteúdo antigo
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if flag:
        run.font.color.rgb = RGBColor(0xA2, 0x2D, 0x2D)
        run.bold = True


def _insert_after_heading(doc: Document, heading_text: str, value: str, flag: bool) -> bool:
    """Insere um parágrafo com o valor preenchido imediatamente a seguir ao heading pedido."""
    heading_norm = _normalize(heading_text)
    for i, paragraph in enumerate(doc.paragraphs):
        if _normalize(paragraph.text) == heading_norm:
            new_para = _insert_paragraph_after(paragraph)
            run = new_para.add_run(value)
            run.font.size = Pt(10)
            if flag:
                run.font.color.rgb = RGBColor(0xA2, 0x2D, 0x2D)
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x2E)
            return True
    return False


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _insert_paragraph_after(paragraph):
    """Insere um novo parágrafo vazio logo após `paragraph` e devolve-o."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    # embrulhar em objecto Paragraph de alto nível
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


# ---------------------------------------------------------------------------
# Upload de template
# ---------------------------------------------------------------------------
def set_custom_template(session_id: str, file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Guarda um template enviado pelo utilizador e associa-o à sessão."""
    session = get_or_create_session(session_id)
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", filename) or "template.docx"
    dest = PMCF_OUTPUT_DIR / f"template_{session.session_id}_{safe_name}"
    with open(dest, "wb") as fh:
        fh.write(file_bytes)

    # validar que é um .docx que o python-docx consegue abrir
    try:
        Document(str(dest))
    except Exception as exc:  # noqa: BLE001
        dest.unlink(missing_ok=True)
        raise ValueError(f"O ficheiro enviado não é um .docx válido: {exc}")

    session.template_path = str(dest)
    return {
        "session_id": session.session_id,
        "template_path": session.template_path,
        "template_name": safe_name,
    }


# ---------------------------------------------------------------------------
# Download
# ---------------------------------------------------------------------------
def get_generated_path(session_id: str) -> Path:
    session = get_session(session_id)
    if not session or not session.generated_path:
        raise FileNotFoundError("Documento ainda não gerado para esta sessão.")
    path = Path(session.generated_path)
    if not path.exists():
        raise FileNotFoundError("Ficheiro gerado já não está disponível no disco.")
    return path


# ---------------------------------------------------------------------------
# Inspeção da sessão (útil para debugging e frontend)
# ---------------------------------------------------------------------------
def session_state(session_id: str) -> Dict[str, Any]:
    session = get_session(session_id)
    if not session:
        raise ValueError("Sessão regulatória não encontrada.")
    return {
        "session_id": session.session_id,
        "step": session.step,
        "collected_keys": sorted(session.collected.keys()),
        "missing_keys": list(session.missing_keys),
        "flagged_fields": list(session.flagged_fields),
        "has_document": bool(session.generated_path),
        "template_path": session.template_path,
    }
