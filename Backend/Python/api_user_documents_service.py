from __future__ import annotations

import io
import os
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

import ollama
from dotenv import load_dotenv

from api_db import db_cursor
from rag_chromadb_service import (
    get_user_documents_collection,
    query_user_documents_chroma,
    delete_user_document_from_chroma,
)

BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent.parent
ENV_PATH = PROJECT_ROOT / "Backend" / ".env"

load_dotenv(dotenv_path=ENV_PATH)

OLLAMA_EMBED_MODEL = os.getenv("OLLAMA_EMBED_MODEL")

USER_DOCS_UPLOAD_DIR = (
    PROJECT_ROOT / os.getenv("USER_DOCS_UPLOAD_DIR", "uploads/user_documents")
).resolve()
USER_DOCS_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_USER_DOCUMENT_SIZE_MB = int(os.getenv("MAX_USER_DOCUMENT_SIZE_MB", "20"))
MAX_USER_DOCUMENT_BYTES = MAX_USER_DOCUMENT_SIZE_MB * 1024 * 1024

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


USER_DOCUMENTS_DDL = """
IF NOT EXISTS (
    SELECT 1
    FROM sys.tables
    WHERE name = 'user_documents'
      AND schema_id = SCHEMA_ID('dbo')
)
BEGIN
    CREATE TABLE dbo.user_documents (
        id                UNIQUEIDENTIFIER NOT NULL PRIMARY KEY,
        user_id           UNIQUEIDENTIFIER NOT NULL,
        original_filename NVARCHAR(255)    NOT NULL,
        stored_filename   NVARCHAR(255)    NOT NULL,
        file_path         NVARCHAR(1000)   NOT NULL,
        mime_type         NVARCHAR(150)    NULL,
        size_bytes        BIGINT           NOT NULL,
        status            NVARCHAR(30)     NOT NULL DEFAULT 'processing',
        error_message     NVARCHAR(MAX)    NULL,
        chunk_count       INT              NOT NULL DEFAULT 0,
        created_at        DATETIME2        NOT NULL DEFAULT SYSUTCDATETIME(),
        updated_at        DATETIME2        NOT NULL DEFAULT SYSUTCDATETIME(),

        CONSTRAINT FK_user_documents_auth_users
            FOREIGN KEY (user_id)
            REFERENCES dbo.auth_users(id)
            ON DELETE CASCADE,

        CONSTRAINT CK_user_documents_status
            CHECK (status IN ('processing', 'ready', 'error'))
    );

    CREATE INDEX IX_user_documents_user_created
        ON dbo.user_documents(user_id, created_at DESC);
END;

IF NOT EXISTS (
    SELECT 1
    FROM sys.tables
    WHERE name = 'user_document_chunks'
      AND schema_id = SCHEMA_ID('dbo')
)
BEGIN
    CREATE TABLE dbo.user_document_chunks (
        id             UNIQUEIDENTIFIER NOT NULL PRIMARY KEY,
        document_id    UNIQUEIDENTIFIER NOT NULL,
        user_id        UNIQUEIDENTIFIER NOT NULL,
        chunk_index    INT              NOT NULL,
        chunk_text     NVARCHAR(MAX)    NOT NULL,
        token_count    INT              NULL,
        page_start     INT              NULL,
        page_end       INT              NULL,
        citation_label NVARCHAR(700)    NOT NULL,
        created_at     DATETIME2        NOT NULL DEFAULT SYSUTCDATETIME(),

        CONSTRAINT FK_user_document_chunks_document
            FOREIGN KEY (document_id)
            REFERENCES dbo.user_documents(id)
            ON DELETE CASCADE,

        CONSTRAINT FK_user_document_chunks_auth_users
            FOREIGN KEY (user_id)
            REFERENCES dbo.auth_users(id)
    );

    CREATE INDEX IX_user_document_chunks_document
        ON dbo.user_document_chunks(document_id, chunk_index);

    CREATE INDEX IX_user_document_chunks_user
        ON dbo.user_document_chunks(user_id);
END;
"""


@dataclass
class TextPage:
    page_number: Optional[int]
    text: str


def init_user_documents_schema() -> None:
    with db_cursor(commit=True) as cur:
        cur.execute(USER_DOCUMENTS_DDL)


def _iso(value: Any) -> Optional[str]:
    if value is None:
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def _safe_filename(filename: str) -> str:
    name = Path(filename or "documento").name
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" ._")
    return name or "documento"


def _row_to_document(row: Any) -> Dict[str, Any]:
    return {
        "id": str(row[0]),
        "user_id": str(row[1]),
        "original_filename": row[2],
        "stored_filename": row[3],
        "file_path": row[4],
        "mime_type": row[5],
        "size_bytes": int(row[6] or 0),
        "status": row[7],
        "error_message": row[8],
        "chunk_count": int(row[9] or 0),
        "created_at": _iso(row[10]),
        "updated_at": _iso(row[11]),
    }


_DOCUMENT_COLUMNS = """
id, user_id, original_filename, stored_filename, file_path, mime_type,
size_bytes, status, error_message, chunk_count, created_at, updated_at
"""


def get_user_document(*, user_id: str, document_id: str) -> Dict[str, Any]:
    with db_cursor() as cur:
        cur.execute(
            f"""
            SELECT {_DOCUMENT_COLUMNS}
              FROM dbo.user_documents
             WHERE id = ? AND user_id = ?
            """,
            document_id,
            user_id,
        )
        row = cur.fetchone()

    if not row:
        raise KeyError("Documento não encontrado.")

    return _row_to_document(row)


def list_user_documents(*, user_id: str, limit: int = 100) -> List[Dict[str, Any]]:
    with db_cursor() as cur:
        cur.execute(
            f"""
            SELECT TOP (?) {_DOCUMENT_COLUMNS}
              FROM dbo.user_documents
             WHERE user_id = ?
             ORDER BY created_at DESC
            """,
            limit,
            user_id,
        )
        rows = cur.fetchall()

    return [_row_to_document(r) for r in rows]


def _extract_txt(content: bytes) -> List[TextPage]:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return [TextPage(page_number=None, text=content.decode(encoding))]
        except UnicodeDecodeError:
            continue
    return [TextPage(page_number=None, text=content.decode("utf-8", errors="ignore"))]


def _extract_docx(content: bytes) -> List[TextPage]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Falta instalar python-docx: pip install python-docx") from exc

    doc = Document(io.BytesIO(content))

    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)

    return [TextPage(page_number=None, text="\n".join(parts))]


def _extract_pdf(content: bytes) -> List[TextPage]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Falta instalar pypdf: pip install pypdf") from exc

    reader = PdfReader(io.BytesIO(content))
    pages: List[TextPage] = []

    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        text = text.strip()
        if text:
            pages.append(TextPage(page_number=idx, text=text))

    return pages


def extract_text_pages(*, filename: str, content: bytes) -> List[TextPage]:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(content)

    if ext == ".docx":
        return _extract_docx(content)

    if ext == ".txt":
        return _extract_txt(content)

    raise ValueError("Formato não suportado. Aceites: PDF, DOCX e TXT.")


def _normalize_text(text: str) -> str:
    text = text or ""
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def chunk_text_pages(
    pages: List[TextPage],
    *,
    max_chars: int = 1800,
    overlap_chars: int = 180,
) -> List[Dict[str, Any]]:
    chunks: List[Dict[str, Any]] = []

    for page in pages:
        text = _normalize_text(page.text)
        if not text:
            continue

        start = 0
        while start < len(text):
            end = min(len(text), start + max_chars)
            chunk = text[start:end].strip()

            if chunk:
                chunks.append({
                    "text": chunk,
                    "page_start": page.page_number,
                    "page_end": page.page_number,
                })

            if end >= len(text):
                break

            start = max(0, end - overlap_chars)

    return chunks


def _token_count(text: str) -> int:
    return len(re.findall(r"\w+", text or ""))


def _update_document_status(
    *,
    document_id: str,
    status: str,
    error_message: Optional[str] = None,
    chunk_count: Optional[int] = None,
) -> None:
    with db_cursor(commit=True) as cur:
        cur.execute(
            """
            UPDATE dbo.user_documents
               SET status = ?,
                   error_message = ?,
                   chunk_count = ISNULL(?, chunk_count),
                   updated_at = SYSUTCDATETIME()
             WHERE id = ?
            """,
            status,
            error_message,
            chunk_count,
            document_id,
        )


def _clear_document_chunks(*, user_id: str, document_id: str) -> None:
    delete_user_document_from_chroma(user_id=user_id, document_id=document_id)

    with db_cursor(commit=True) as cur:
        cur.execute(
            "DELETE FROM dbo.user_document_chunks WHERE document_id = ? AND user_id = ?",
            document_id,
            user_id,
        )


def _insert_chunks_sql(
    *,
    user_id: str,
    document_id: str,
    original_filename: str,
    chunks: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    inserted: List[Dict[str, Any]] = []

    with db_cursor(commit=True) as cur:
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            page_start = chunk.get("page_start")
            page_end = chunk.get("page_end")

            if page_start:
                citation = f"Documento interno: {original_filename} pp. {page_start}"
                if page_end and page_end != page_start:
                    citation = f"Documento interno: {original_filename} pp. {page_start}-{page_end}"
            else:
                citation = f"Documento interno: {original_filename} chunk {i + 1}"

            text = chunk["text"]

            cur.execute(
                """
                INSERT INTO dbo.user_document_chunks (
                    id, document_id, user_id, chunk_index, chunk_text,
                    token_count, page_start, page_end, citation_label
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                chunk_id,
                document_id,
                user_id,
                i,
                text,
                _token_count(text),
                page_start,
                page_end,
                citation,
            )

            inserted.append({
                "id": chunk_id,
                "document_id": document_id,
                "user_id": user_id,
                "chunk_index": i,
                "chunk_text": text,
                "token_count": _token_count(text),
                "page_start": page_start,
                "page_end": page_end,
                "citation_label": citation,
                "original_filename": original_filename,
            })

    return inserted


def _embed_and_index_chunks(
    *,
    chunks: List[Dict[str, Any]],
    user_id: str,
    document_id: str,
    original_filename: str,
) -> None:
    if not chunks:
        return

    if not OLLAMA_EMBED_MODEL:
        raise ValueError("Falta OLLAMA_EMBED_MODEL no .env")

    collection = get_user_documents_collection()

    batch_size = 32
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start:start + batch_size]
        texts = [c["chunk_text"] for c in batch]

        embed_result = ollama.embed(
            model=OLLAMA_EMBED_MODEL,
            input=texts,
        )
        embeddings = embed_result["embeddings"]

        ids = [str(c["id"]) for c in batch]
        documents = texts
        metadatas = []

        for c in batch:
            metadatas.append({
                "source_type": "user_document",
                "user_id": str(user_id),
                "document_id": str(document_id),
                "chunk_id": str(c["id"]),
                "original_filename": str(original_filename),
                "chunk_index": int(c["chunk_index"]),
                "page_start": int(c["page_start"]) if c["page_start"] is not None else -1,
                "page_end": int(c["page_end"]) if c["page_end"] is not None else -1,
                "citation_label": str(c["citation_label"]),
                "short_name": "USER_DOC",
                "section_type": "user_document",
                "section_number": "",
                "section_title": str(original_filename),
            })

        collection.upsert(
            ids=ids,
            documents=documents,
            metadatas=metadatas,
            embeddings=embeddings,
        )


def process_user_document(
    *,
    user_id: str,
    document_id: str,
    content: bytes,
    original_filename: str,
) -> Dict[str, Any]:
    try:
        _clear_document_chunks(user_id=user_id, document_id=document_id)

        pages = extract_text_pages(filename=original_filename, content=content)
        chunks = chunk_text_pages(pages)

        if not chunks:
            raise ValueError("Não foi possível extrair texto útil do documento.")

        inserted_chunks = _insert_chunks_sql(
            user_id=user_id,
            document_id=document_id,
            original_filename=original_filename,
            chunks=chunks,
        )

        _embed_and_index_chunks(
            chunks=inserted_chunks,
            user_id=user_id,
            document_id=document_id,
            original_filename=original_filename,
        )

        _update_document_status(
            document_id=document_id,
            status="ready",
            error_message=None,
            chunk_count=len(inserted_chunks),
        )

    except Exception as exc:
        _update_document_status(
            document_id=document_id,
            status="error",
            error_message=str(exc),
            chunk_count=0,
        )

    return get_user_document(user_id=user_id, document_id=document_id)


def upload_user_document(
    *,
    user_id: str,
    original_filename: str,
    content: bytes,
    mime_type: Optional[str],
) -> Dict[str, Any]:
    if not content:
        raise ValueError("O ficheiro está vazio.")

    if len(content) > MAX_USER_DOCUMENT_BYTES:
        raise ValueError(f"O ficheiro excede o limite de {MAX_USER_DOCUMENT_SIZE_MB} MB.")

    safe_name = _safe_filename(original_filename)
    ext = Path(safe_name).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Formato não suportado. Aceites: PDF, DOCX e TXT.")

    document_id = str(uuid.uuid4())
    stored_filename = f"userdoc_{user_id}_{document_id}_{safe_name}"
    user_dir = USER_DOCS_UPLOAD_DIR / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)

    file_path = user_dir / stored_filename
    file_path.write_bytes(content)

    now = datetime.now(timezone.utc)

    with db_cursor(commit=True) as cur:
        cur.execute(
            """
            INSERT INTO dbo.user_documents (
                id, user_id, original_filename, stored_filename, file_path,
                mime_type, size_bytes, status, error_message, chunk_count,
                created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, 'processing', NULL, 0, ?, ?)
            """,
            document_id,
            user_id,
            safe_name,
            stored_filename,
            str(file_path),
            (mime_type or "application/octet-stream")[:150],
            len(content),
            now,
            now,
        )

    return process_user_document(
        user_id=user_id,
        document_id=document_id,
        content=content,
        original_filename=safe_name,
    )


def delete_user_document(
    *,
    user_id: str,
    document_id: str,
) -> None:
    doc = get_user_document(user_id=user_id, document_id=document_id)

    delete_user_document_from_chroma(user_id=user_id, document_id=document_id)

    path = Path(doc["file_path"])
    if path.exists():
        try:
            path.unlink()
        except Exception:
            pass

    with db_cursor(commit=True) as cur:
        cur.execute(
            "DELETE FROM dbo.user_documents WHERE id = ? AND user_id = ?",
            document_id,
            user_id,
        )

        if cur.rowcount == 0:
            raise KeyError("Documento não encontrado.")


def reprocess_user_document(
    *,
    user_id: str,
    document_id: str,
) -> Dict[str, Any]:
    doc = get_user_document(user_id=user_id, document_id=document_id)
    path = Path(doc["file_path"])

    if not path.exists():
        raise FileNotFoundError("Ficheiro original não encontrado no disco.")

    _update_document_status(
        document_id=document_id,
        status="processing",
        error_message=None,
        chunk_count=0,
    )

    return process_user_document(
        user_id=user_id,
        document_id=document_id,
        content=path.read_bytes(),
        original_filename=doc["original_filename"],
    )


def query_user_documents_for_rag(
    *,
    user_id: str,
    question: str,
    n_results: int = 4,
) -> Dict[str, Any]:
    if not question or not question.strip():
        return {
            "records": [],
            "scores": [],
        }

    if not OLLAMA_EMBED_MODEL:
        return {
            "records": [],
            "scores": [],
        }

    try:
        query_embedding = ollama.embed(
            model=OLLAMA_EMBED_MODEL,
            input=question,
        )["embeddings"][0]

        result = query_user_documents_chroma(
            query_embedding=query_embedding,
            n_results=n_results,
            where={"user_id": str(user_id)},
        )
    except Exception:
        return {
            "records": [],
            "scores": [],
        }

    metadatas = (result.get("metadatas") or [[]])[0]
    documents = (result.get("documents") or [[]])[0]
    distances = (result.get("distances") or [[]])[0]
    ids = (result.get("ids") or [[]])[0]

    records: List[Dict[str, Any]] = []
    scores: List[float] = []

    def clean_page(value: Any) -> Optional[int]:
        try:
            if value in (None, "", -1, "-1"):
                return None
            return int(value)
        except Exception:
            return None

    for i, meta in enumerate(metadatas):
        meta = meta or {}
        distance = float(distances[i]) if i < len(distances) else 999.0
        score = 1.0 / (1.0 + max(0.0, distance))

        # Evita meter contexto fraco só porque existe documento carregado.
        if score < 0.25:
            continue

        text = documents[i] if i < len(documents) else ""

        record = {
            "chunk_id": ids[i] if i < len(ids) else meta.get("chunk_id", i),
            "citation_label": str(meta.get("citation_label", "") or ""),
            "short_name": "USER_DOC",
            "section_type": "user_document",
            "section_number": "",
            "section_title": str(meta.get("original_filename", "") or "Documento interno"),
            "page_start": clean_page(meta.get("page_start")),
            "page_end": clean_page(meta.get("page_end")),
            "chunk_text": text,
            "text": text,
            "source_type": "user_document",
            "document_id": str(meta.get("document_id", "") or ""),
            "original_filename": str(meta.get("original_filename", "") or ""),
        }

        records.append(record)
        scores.append(float(score))

    return {
        "records": records,
        "scores": scores,
    }