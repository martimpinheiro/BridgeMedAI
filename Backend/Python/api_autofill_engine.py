"""
Auto-fill Engine do BridgeMedAI.

Camada que pega num template do registry, cruza os `auto_fillable_fields` e
`human_required_fields` declarados pelo template com os `extracted_fields`
persistidos no Context Memory do utilizador, e gera um .docx pré-preenchido:

1. **Substituição de placeholders**: procura `{{field_key}}` em parágrafos
   (incluindo dentro de tabelas) e substitui pelo valor conhecido. Se o campo
   estiver vazio, marca com um sinal visual de revisão manual.
2. **Cover sheet de contexto**: insere uma página inicial "Contexto
   auto-preenchido" listando todos os campos conhecidos do perfil — útil mesmo
   para templates que não tenham placeholders explícitos.
3. **Cálculo de cobertura + lifecycle**: decide automaticamente se a
   instância deve passar a `partial`, `awaiting` ou ficar como está.

Templates .xlsx não são suportados nesta iteração — devolvem erro claro;
o estado da instância fica intocado para que o utilizador trate manualmente.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

from api_template_registry import TemplateRecord, get_record, get_template_file_path
from api_context_memory import (
    get_instance,
    get_profile,
    list_fields,
    list_instances,
    update_instance,
)

# ---------------------------------------------------------------------------
# Output dir (partilhado com o regulatory service)
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parents[2]
AUTOFILL_OUTPUT_DIR = (
    PROJECT_ROOT / os.getenv("PMCF_OUTPUT_DIR", "Backend/generated")
) / "copilot"
AUTOFILL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-.]+)\s*\}\}")
MANUAL_FLAG_PREFIX = "⚠️"


# ---------------------------------------------------------------------------
# Coverage + lifecycle
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class Coverage:
    total: int
    filled: int
    missing: List[str]
    missing_human_required: List[str]
    missing_auto_fillable: List[str]
    coverage_pct: int

    def to_dict(self) -> Dict[str, Any]:
        return {
            "total": self.total,
            "filled": self.filled,
            "coverage_pct": self.coverage_pct,
            "missing": self.missing,
            "missing_human_required": self.missing_human_required,
            "missing_auto_fillable": self.missing_auto_fillable,
        }


def compute_coverage(
    record: TemplateRecord,
    fields_by_key: Dict[str, str],
) -> Coverage:
    required_auto = list(record.auto_fillable_fields)
    required_human = list(record.human_required_fields)
    seen: Set[str] = set()
    required: List[Tuple[str, str]] = []  # (key, kind)
    for k in required_auto:
        if k not in seen:
            seen.add(k)
            required.append((k, "auto_fillable"))
    for k in required_human:
        if k not in seen:
            seen.add(k)
            required.append((k, "human_required"))

    missing: List[str] = []
    missing_human: List[str] = []
    missing_auto: List[str] = []
    filled = 0
    for key, kind in required:
        value = (fields_by_key.get(key) or "").strip()
        if value:
            filled += 1
        else:
            missing.append(key)
            if kind == "human_required":
                missing_human.append(key)
            else:
                missing_auto.append(key)

    total = len(required)
    pct = int(round((filled / total) * 100)) if total else 100

    return Coverage(
        total=total,
        filled=filled,
        missing=missing,
        missing_human_required=missing_human,
        missing_auto_fillable=missing_auto,
        coverage_pct=pct,
    )


_AUTO_PROMOTABLE_STATES = {"draft", "partial", "awaiting"}


def decide_state(coverage: Coverage, current_state: str) -> str:
    """Apenas promove/demove entre draft / partial / awaiting.
    Estados explicitamente humanos (reviewed/approved/exported) nunca são
    alterados automaticamente."""
    if current_state not in _AUTO_PROMOTABLE_STATES:
        return current_state

    if coverage.total == 0:
        return current_state

    if coverage.filled == 0:
        return "draft"

    if coverage.coverage_pct >= 80 and coverage.missing_human_required:
        return "awaiting"

    if coverage.coverage_pct >= 100:
        return "awaiting" if coverage.missing_human_required else "awaiting"

    return "partial"


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------
def _styled_run(paragraph, text: str, *, flagged: bool) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    if flagged:
        run.font.color.rgb = RGBColor(0xA2, 0x2D, 0x2D)
        run.bold = True
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x2E)


def _replace_in_paragraph(
    paragraph,
    substitutions: Dict[str, str],
) -> int:
    """Substitui ocorrências literais de `<...>` no parágrafo pelos valores
    do mapping `substitutions` (chaves no formato '<texto>'). Devolve número
    de substituições efetivas neste parágrafo.

    Como o python-docx parte o texto em runs com estilos diferentes (e os
    placeholders Fraunhofer estão geralmente num único run amarelo/itálico),
    fazemos substituição agregada: juntamos o texto do parágrafo, substituímos,
    e reinjetamos como um único run preservando o estilo do parágrafo.
    """
    full_text = paragraph.text or ""
    if not full_text or "<" not in full_text:
        return 0

    new_text = full_text
    count = 0
    for full_match, value in substitutions.items():
        if full_match in new_text:
            new_text = new_text.replace(full_match, value or "")
            count += 1

    if count == 0:
        return 0

    # Limpar runs existentes preservando estilo do parágrafo, e reinjetar
    for run in list(paragraph.runs):
        run.text = ""
    _styled_run(paragraph, new_text, flagged=False)
    return count


def _replace_in_doc(doc, substitutions: Dict[str, str]) -> Dict[str, Any]:
    """Substitui em paragraphs (body), tabelas e headers/footers."""
    total = 0

    for p in doc.paragraphs:
        total += _replace_in_paragraph(p, substitutions)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += _replace_in_paragraph(p, substitutions)

    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                total += _replace_in_paragraph(p, substitutions)
            for t in hf.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            total += _replace_in_paragraph(p, substitutions)

    return {
        "substituted_count": total,
        "substitution_keys": list(substitutions.keys()),
    }


def _build_cover_sheet(
    doc,
    *,
    record: TemplateRecord,
    profile: Dict[str, Any],
    substituted_count: int,
    total_placeholders: int,
) -> None:
    """Cabeçalho mínimo no topo do documento — só identifica a geração e
    quantos placeholders foram substituídos. Os valores já estão no corpo."""
    body = doc.element.body
    insertion_point = body[0] if len(body) > 0 else None

    when = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    paragraphs_to_add: List[Tuple[str, str]] = [
        ("title", f"{record.id} — {record.name}"),
        (
            "subtitle",
            "Documento parcialmente auto-preenchido pelo Regulatory Documentation Copilot do BridgeMedAI.",
        ),
        (
            "meta",
            " · ".join([
                f"Produto: {profile.get('name') or '(sem nome)'}",
                f"Gerado em: {when}",
                f"Substituições aplicadas: {substituted_count}/{total_placeholders}",
            ]),
        ),
        ("rule", "—" * 24),
    ]

    # Cada chamada a `insertion_point.addprevious(new)` adiciona `new`
    # imediatamente antes de `insertion_point`, mas DEPOIS de qualquer
    # `new` inserido anteriormente. Logo, a ordem visual final é a mesma
    # da ordem de iteração — iteramos em ordem natural.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph as _Paragraph

    for kind, text in paragraphs_to_add:
        new_p_elem = OxmlElement("w:p")
        if insertion_point is not None:
            insertion_point.addprevious(new_p_elem)
        else:
            body.append(new_p_elem)
        paragraph = _Paragraph(new_p_elem, doc)
        run = paragraph.add_run(text)
        if kind == "title":
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = RGBColor(0x15, 0x2A, 0x20)
        elif kind == "subtitle":
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2A, 0x4F, 0x3E)
        elif kind == "meta":
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x50)
        elif kind == "section":
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x2E)
        elif kind == "field_filled":
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x2E)
        elif kind == "field_missing":
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA2, 0x2D, 0x2D)
            run.bold = True
        elif kind == "note":
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x50)
        elif kind == "rule":
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xB8, 0xCD, 0xA8)


def render_template_docx(
    *,
    template_path: Path,
    record: TemplateRecord,
    profile: Dict[str, Any],
    placeholder_substitutions: Optional[Dict[str, str]] = None,
    total_placeholders: int = 0,
    output_path: Path,
) -> Dict[str, Any]:
    """Gera o .docx final substituindo placeholders `<...>` no corpo + adiciona
    cover sheet mínimo. Se `placeholder_substitutions` for None ou vazio,
    o documento sai praticamente igual ao template (só com o cabeçalho)."""
    doc = Document(str(template_path))
    subs = placeholder_substitutions or {}
    report = _replace_in_doc(doc, subs)
    _build_cover_sheet(
        doc,
        record=record,
        profile=profile,
        substituted_count=report["substituted_count"],
        total_placeholders=total_placeholders or len(subs),
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return report


# ---------------------------------------------------------------------------
# Orquestração por instância
# ---------------------------------------------------------------------------
class AutofillError(RuntimeError):
    pass


def _slugify(text: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", (text or "").strip())
    return safe[:60] or "untitled"


def _fields_by_key(profile_id: str, user_id: str) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for f in list_fields(profile_id=profile_id, user_id=user_id):
        if f.get("field_value"):
            out[f["field_key"]] = f["field_value"]
    return out


def autofill_instance(
    *,
    instance_id: str,
    user_id: str,
    placeholder_substitutions: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    """Gera o `.docx` final para esta instância.

    Se `placeholder_substitutions` for fornecido (vindo do questionário Q&A),
    substitui literalmente cada chave (formato `<texto>`) no corpo do template
    pelo valor correspondente. Caso contrário, gera o template sem alterações
    no corpo (apenas o cabeçalho de identificação).
    """
    instance = get_instance(instance_id=instance_id, user_id=user_id)
    template_id = instance["template_id"]

    try:
        record = get_record(template_id)
    except KeyError as exc:
        raise AutofillError(f"Template '{template_id}' não existe.") from exc

    try:
        template_path = get_template_file_path(template_id)
    except Exception as exc:
        raise AutofillError(f"Ficheiro do template não encontrado: {exc}") from exc

    if template_path.suffix.lower() != ".docx":
        raise AutofillError(
            f"Auto-fill ainda só suporta templates .docx (este é {template_path.suffix}). "
            "Usa o template original para preencher manualmente."
        )

    profile_id = instance["product_profile_id"]
    profile = get_profile(profile_id=profile_id, user_id=user_id)
    fields_by_key = _fields_by_key(profile_id, user_id)
    coverage = compute_coverage(record, fields_by_key)

    # Total de placeholders no template (para mostrar X/Y no cover sheet)
    try:
        from api_template_placeholders import extract_placeholders
        all_phs = extract_placeholders(template_path)
        total_placeholders = len(all_phs)
    except Exception:
        total_placeholders = 0

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"{record.id}_{_slugify(record.name)}_{timestamp}.docx"
    out_path = AUTOFILL_OUTPUT_DIR / instance_id / out_name

    replacement_report = render_template_docx(
        template_path=template_path,
        record=record,
        profile=profile,
        placeholder_substitutions=placeholder_substitutions,
        total_placeholders=total_placeholders,
        output_path=out_path,
    )

    new_state = decide_state(coverage, instance["state"])
    update_payload: Dict[str, Any] = {
        "file_path": str(out_path),
        "download_name": out_name,
    }
    if new_state != instance["state"]:
        update_payload["state"] = new_state

    updated_instance = update_instance(
        instance_id=instance_id,
        user_id=user_id,
        **update_payload,
    )

    return {
        "instance": updated_instance,
        "template": {"id": record.id, "name": record.name, "category": record.category},
        "coverage": coverage.to_dict(),
        "replacement_report": replacement_report,
        "new_state": new_state,
        "previous_state": instance["state"],
    }


def autofill_all_for_profile(
    *,
    profile_id: str,
    user_id: str,
) -> List[Dict[str, Any]]:
    instances = list_instances(profile_id=profile_id, user_id=user_id)
    out: List[Dict[str, Any]] = []
    for inst in instances:
        try:
            result = autofill_instance(instance_id=inst["id"], user_id=user_id)
            out.append({"ok": True, "instance_id": inst["id"], **result})
        except AutofillError as exc:
            out.append({"ok": False, "instance_id": inst["id"], "error": str(exc)})
    return out


def get_generated_file_path(*, instance_id: str, user_id: str) -> Path:
    instance = get_instance(instance_id=instance_id, user_id=user_id)
    path_str = instance.get("file_path")
    if not path_str:
        raise FileNotFoundError("Esta instância ainda não foi auto-preenchida.")
    path = Path(path_str)
    if not path.exists():
        raise FileNotFoundError(f"Ficheiro não existe em disco: {path}")

    # Defensive: garantir que o caminho está dentro do diretório de output
    try:
        path.resolve().relative_to(AUTOFILL_OUTPUT_DIR.resolve())
    except ValueError as exc:
        raise FileNotFoundError("Ficheiro fora do diretório permitido.") from exc

    return path
