"""
Extrator de placeholders `<...>` reais nos templates Fraunhofer.

Os templates Fraunhofer usam marcadores no formato `<descrição/instrução>`
nas zonas que o autor humano deve preencher (nome do produto, versão, etc.).
Este módulo:

1. Abre o .docx
2. Encontra todos os `<...>` em paragraphs (incluindo dentro de tabelas
   e header/footer)
3. Devolve uma lista ordenada de placeholders com texto + contexto + chave
   estável (hash) — para serem perguntados em chat e depois substituídos

Não é tentativa de "ler entre linhas" das instruções italic/amarelas dos
templates. Essas continuam por preencher manualmente — é por design dos
Fraunhofer (cada secção tem orientação para o autor humano).
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document

# Captura <texto> com no máximo 500 chars internos, sem `<` ou `>` aninhados.
# Permite quebras de linha (Word às vezes parte texto longo em runs).
_PLACEHOLDER_RE = re.compile(r"<([^<>]{1,500})>", re.DOTALL)


@dataclass(frozen=True)
class Placeholder:
    """Um marcador `<...>` encontrado num template Fraunhofer."""

    key: str            # hash curto para identificação estável
    raw_text: str       # texto exato dentro de `<>` (sem os angle brackets)
    full_match: str     # com angle brackets, para substituição: "<...>"
    context: str        # cabeçalho/secção próxima — ajuda a perceber o que pedir
    location: str       # ex: "body", "t2r1c0", "header"
    order: int          # ordem em que aparece no documento

    def to_dict(self) -> Dict[str, Any]:
        return {
            "key": self.key,
            "raw_text": self.raw_text,
            "full_match": self.full_match,
            "context": self.context,
            "location": self.location,
            "order": self.order,
        }


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def _make_key(raw_text: str, location: str) -> str:
    """Chave estável para o placeholder. Inclui location para distinguir
    placeholders idênticos em sítios diferentes (raro mas possível)."""
    h = hashlib.sha1(f"{location}::{_normalize(raw_text).lower()}".encode("utf-8")).hexdigest()
    return h[:10]


def _placeholders_in_text(
    text: str,
    location: str,
    context: str,
    start_order: int,
) -> List[Placeholder]:
    out: List[Placeholder] = []
    if not text or "<" not in text:
        return out
    for i, match in enumerate(_PLACEHOLDER_RE.finditer(text)):
        raw_inner = match.group(1)
        normalized_inner = _normalize(raw_inner)
        # Filtra falsos positivos óbvios (tags HTML/XML que escaparam, datas
        # tipo <1, comparadores). Heurística: aceitar se tiver pelo menos
        # uma letra alfabética.
        if not re.search(r"[A-Za-zÀ-ÿ]", normalized_inner):
            continue
        # Filtra também tags de markup conhecidas
        if normalized_inner.lower() in {"br", "hr", "p", "b", "i", "u"}:
            continue
        full = f"<{raw_inner}>"
        loc_with_idx = f"{location}#{i}" if i else location
        out.append(
            Placeholder(
                key=_make_key(raw_inner, loc_with_idx),
                raw_text=normalized_inner,
                full_match=full,
                context=context,
                location=loc_with_idx,
                order=start_order + i,
            )
        )
    return out


def _nearest_heading_for_paragraph(
    paragraphs: List[Any],
    index: int,
    max_lookback: int = 10,
) -> str:
    """Procura para trás até `max_lookback` parágrafos por algo que pareça
    cabeçalho (estilo Heading X ou texto curto e isolado)."""
    for back in range(index - 1, max(-1, index - max_lookback), -1):
        p = paragraphs[back]
        style = getattr(getattr(p, "style", None), "name", "") or ""
        if "Heading" in style or "Title" in style:
            return _normalize(p.text)
        # heurística: parágrafo curto não-vazio sem ponto final
        txt = _normalize(p.text)
        if 3 < len(txt) < 80 and not txt.endswith(".") and not txt.endswith(":"):
            # provável cabeçalho informal
            return txt
    return ""


def extract_placeholders(template_path: Path) -> List[Placeholder]:
    """Devolve todos os placeholders `<...>` do template, com contexto."""
    doc = Document(str(template_path))
    out: List[Placeholder] = []
    order = 0

    # 1) Parágrafos do body com contexto = cabeçalho mais próximo
    paragraphs = list(doc.paragraphs)
    for pi, p in enumerate(paragraphs):
        ctx = _nearest_heading_for_paragraph(paragraphs, pi)
        found = _placeholders_in_text(p.text, "body", ctx, order)
        out.extend(found)
        order += len(found)

    # 2) Tabelas: contexto = cabeçalho da coluna (primeira linha) se existir
    for ti, table in enumerate(doc.tables):
        header_row = table.rows[0] if table.rows else None
        col_headers: List[str] = []
        if header_row:
            for cell in header_row.cells:
                txt = ""
                if cell.paragraphs:
                    txt = _normalize(cell.paragraphs[0].text)
                col_headers.append(txt)

        for ri, row in enumerate(table.rows):
            if ri == 0 and header_row is not None:
                # Saltamos a header row (não costuma ter placeholders relevantes)
                # mas processamos caso o user tenha posto algo lá
                pass
            for ci, cell in enumerate(row.cells):
                col_header = col_headers[ci] if ci < len(col_headers) else ""
                for p in cell.paragraphs:
                    loc = f"t{ti}r{ri}c{ci}"
                    ctx = col_header or f"tabela {ti}"
                    found = _placeholders_in_text(p.text, loc, ctx, order)
                    out.extend(found)
                    order += len(found)

    # 3) Header / footer
    for sec_i, sec in enumerate(doc.sections):
        for hf_label, hf in [("header", sec.header), ("footer", sec.footer)]:
            if hf is None:
                continue
            for p in hf.paragraphs:
                loc = f"{hf_label}{sec_i}"
                found = _placeholders_in_text(p.text, loc, hf_label, order)
                out.extend(found)
                order += len(found)
            for ti, t in enumerate(hf.tables):
                for ri, row in enumerate(t.rows):
                    for ci, cell in enumerate(row.cells):
                        for p in cell.paragraphs:
                            loc = f"{hf_label}{sec_i}-t{ti}r{ri}c{ci}"
                            found = _placeholders_in_text(p.text, loc, hf_label, order)
                            out.extend(found)
                            order += len(found)

    # Dedup mantendo a primeira ocorrência (chaves iguais = pergunta única)
    seen: Set[str] = set()
    deduped: List[Placeholder] = []
    for ph in out:
        if ph.key in seen:
            continue
        seen.add(ph.key)
        deduped.append(ph)

    # Ordenar por ordem de aparecimento
    deduped.sort(key=lambda p: p.order)
    return deduped


# ---------------------------------------------------------------------------
# Construção de perguntas em PT para um placeholder
# ---------------------------------------------------------------------------

def humanize_placeholder_text(raw_text: str, context: str = "") -> str:
    """
    Converte o texto cru de um placeholder do template numa pergunta clara
    em português de Portugal.

    A função é centralizada para servir todos os templates .docx:
    Clinical Evaluation, Software, Risk Management, Usability, Vigilance, PMS/PMCF, etc.
    """
    raw = _normalize(raw_text or "")
    t = raw.lower()
    ctx = _normalize(context or "").lower()

    def has_any(*terms: str) -> bool:
        return any(term in t for term in terms)

    # ---------------------------------------------------------
    # Casos numéricos genéricos
    # Ex.: "enter number between 0 and 3"
    # ---------------------------------------------------------
    m = re.search(r"enter\s+number\s+between\s+(\d+)\s+and\s+(\d+)", t)
    if m:
        low, high = m.group(1), m.group(2)

        if "clinical equivalence" in ctx or "equivalence" in ctx:
            return (
                f"Quantos dispositivos equivalentes queres considerar na análise de equivalência clínica? "
                f"Indica um número entre {low} e {high}."
            )

        if (
            "post-market" in ctx
            or "surveillance" in ctx
            or "clinical follow-up" in ctx
            or "pmcf" in ctx
            or "pms" in ctx
        ):
            return (
                f"Quantas fontes, atividades ou evidências de PMS/PMCF queres indicar nesta secção? "
                f"Indica um número entre {low} e {high}."
            )

        if "risk" in ctx or "risco" in ctx:
            return f"Que valor numérico de risco queres indicar? Usa um número entre {low} e {high}."

        if "usability" in ctx or "human factors" in ctx:
            return f"Quantos cenários/tarefas de usabilidade queres indicar? Usa um número entre {low} e {high}."

        return f"Indica um número entre {low} e {high} para este campo."

    # ---------------------------------------------------------
    # Mapeamentos exatos comuns
    # ---------------------------------------------------------
    direct = {
        "name of the product": "Qual é o nome comercial do produto?",
        "product name": "Qual é o nome comercial do produto?",
        "device name": "Qual é o nome do dispositivo?",
        "name of device": "Qual é o nome do dispositivo?",

        "version of the product": "Qual é a versão do produto?",
        "version of the software": "Qual é a versão do software?",
        "software version": "Qual é a versão do software?",
        "basic udi-di, if/when available": "Qual é o Basic UDI-DI? Se ainda não existir, escreve “saltar”.",

        "company": "Qual é o nome da empresa/fabricante?",
        "manufacturer": "Qual é o nome do fabricante?",
        "manufacturer name": "Qual é o nome do fabricante?",
        "manufacturer address": "Qual é a morada do fabricante?",
        "address": "Qual é a morada?",
        "email": "Qual é o email de contacto?",
        "phone": "Qual é o telefone de contacto?",

        "name": "Qual é o nome?",
        "date": "Que data devo colocar? Usa o formato AAAA-MM-DD.",
        "author": "Quem é o autor do documento?",
        "reviewer": "Quem é o revisor do documento?",
        "approver": "Quem aprova o documento?",
        "approved by": "Quem aprovou?",
        "created by": "Quem criou?",
        "description of change": "Qual é a descrição da alteração?",
        "version": "Qual é a versão?",
    }

    if t in direct:
        return direct[t]

    # ---------------------------------------------------------
    # Casos longos / frases específicas dos templates
    # ---------------------------------------------------------
    if "version of the product" in t and "standalone software" in t:
        return (
            "Qual é a versão do produto? "
            "Se for software autónomo e só fizer sentido indicar a versão de software, podes escrever “saltar”."
        )

    if "basic udi-di" in t:
        return "Qual é o Basic UDI-DI? Se ainda não estiver disponível, escreve “saltar”."

    if "single registration number" in t or t == "srn" or " srn" in t:
        return "Qual é o Single Registration Number (SRN) do fabricante? Se ainda não existir, escreve “saltar”."

    if "notified body" in t:
        return "Qual é o organismo notificado aplicável? Se não for aplicável, escreve “não aplicável”."

    if "certificate" in t and "number" in t:
        return "Qual é o número do certificado CE? Se ainda não existir ou não for aplicável, escreve “saltar”."

    # ---------------------------------------------------------
    # Clinical Evaluation / PMCF / PMS
    # ---------------------------------------------------------
    if has_any("intended purpose", "intended use"):
        return "Qual é a finalidade prevista do dispositivo, exatamente como deve constar na documentação técnica?"

    if has_any("intended users", "user group", "users"):
        return "Quem são os utilizadores previstos do dispositivo?"

    if has_any("patient population", "target population", "population"):
        return "Qual é a população de doentes/utilizadores alvo?"

    if has_any("medical indication", "indication", "clinical indication"):
        return "Quais são as indicações clínicas previstas?"

    if "contraindication" in t:
        return "Quais são as contraindicações conhecidas? Se não houver, escreve “nenhuma identificada”."

    if "clinical benefit" in t:
        return "Qual é o benefício clínico esperado do dispositivo?"

    if "clinical claim" in t or "claims" in t:
        return "Que alegações clínicas/desempenho queres declarar para o dispositivo?"

    if "clinical evidence" in t:
        return "Que evidência clínica suporta a segurança e o desempenho do dispositivo?"

    if "clinical equivalence" in t:
        return "Que dispositivo(s) equivalente(s), se existirem, devem ser considerados na avaliação clínica?"

    if "pmcf objective" in t or "objectives" in t and ("pmcf" in ctx or "clinical follow-up" in ctx):
        return "Quais são os objetivos específicos do PMCF?"

    if "pmcf method" in t or "methods" in t and ("pmcf" in ctx or "clinical follow-up" in ctx):
        return "Que métodos PMCF queres usar? Por exemplo, literatura, inquéritos, registos, reclamações ou estudo PMCF."

    if "post-market surveillance" in t or "pms" in t:
        return "Que atividades de vigilância pós-comercialização/PMS queres indicar?"

    # ---------------------------------------------------------
    # Software / IEC 62304
    # ---------------------------------------------------------
    if "software safety class" in t or "software safety classification" in t:
        return "Qual é a classe de segurança do software segundo IEC 62304? Indica A, B ou C e a justificação."

    if "software development plan" in t:
        return "Qual é o plano ou referência do plano de desenvolvimento de software?"

    if "software architecture" in t:
        return "Descreve resumidamente a arquitetura do software ou indica o documento onde está descrita."

    if "software requirement" in t or "requirements specification" in t:
        return "Que requisito de software deve ser indicado neste campo?"

    if "system requirement" in t or "stakeholder requirement" in t:
        return "Que requisito de sistema/stakeholder deve ser indicado neste campo?"

    if "test protocol" in t:
        return "Qual é o protocolo de teste aplicável?"

    if "test result" in t:
        return "Qual foi o resultado do teste?"

    if "known anomaly" in t or "known anomalies" in t:
        return "Existem anomalias conhecidas nesta versão? Se não houver, escreve “nenhuma conhecida”."

    if "release" in t and "version" in t:
        return "Qual é a versão/release de software a documentar?"

    # ---------------------------------------------------------
    # Risk Management / ISO 14971
    # ---------------------------------------------------------
    if "hazardous situation" in t:
        return "Qual é a situação perigosa identificada?"

    if "hazard" in t:
        return "Qual é o perigo identificado?"

    if "harm" in t:
        return "Qual é o dano possível para o doente/utilizador?"

    if "severity" in t:
        return "Qual é a severidade estimada do dano?"

    if "probability" in t or "occurrence" in t:
        return "Qual é a probabilidade estimada de ocorrência?"

    if "risk control" in t or "mitigation" in t:
        return "Que medida de controlo/mitigação de risco deve ser indicada?"

    if "residual risk" in t:
        return "Qual é o risco residual após as medidas de controlo?"

    if "risk acceptability" in t or "acceptable" in t:
        return "O risco residual é aceitável? Indica a justificação."

    # ---------------------------------------------------------
    # Usability / Human Factors
    # ---------------------------------------------------------
    if "use scenario" in t:
        return "Qual é o cenário de utilização a considerar?"

    if "critical task" in t:
        return "Qual é a tarefa crítica de utilização?"

    if "use error" in t:
        return "Que erro de utilização pode ocorrer?"

    if "user interface" in t:
        return "Que elemento da interface de utilizador está a ser avaliado?"

    if "formative" in t:
        return "Que atividade/teste formativo de usabilidade queres indicar?"

    if "summative" in t:
        return "Que teste somativo de usabilidade queres indicar?"

    # ---------------------------------------------------------
    # Vigilance / Incident / CAPA
    # ---------------------------------------------------------
    if "incident description" in t or "description of incident" in t:
        return "Descreve o incidente ocorrido."

    if "incident date" in t:
        return "Qual foi a data do incidente? Usa o formato AAAA-MM-DD."

    if "serious incident" in t:
        return "O incidente é grave? Indica sim/não e justifica."

    if "field safety corrective action" in t or "fsca" in t:
        return "Que ação corretiva de segurança no terreno (FSCA) foi ou será tomada?"

    if "field safety notice" in t or "fsn" in t:
        return "Que mensagem principal deve constar no Field Safety Notice?"

    if "competent authority" in t:
        return "Qual é a autoridade competente a notificar?"

    if "root cause" in t:
        return "Qual é a causa raiz identificada?"

    if "corrective action" in t:
        return "Que ação corretiva deve ser registada?"

    if "preventive action" in t:
        return "Que ação preventiva deve ser registada?"

    # ---------------------------------------------------------
    # Datas / versões / pessoas — heurísticas finais
    # ---------------------------------------------------------
    if "date" in t:
        return f"Que data devo colocar para “{raw}”? Usa o formato AAAA-MM-DD."

    if "version" in t:
        return f"Qual é a versão aplicável para “{raw}”?"

    if "name" in t:
        return f"Qual é o nome a preencher neste campo?"

    if "description" in t:
        return f"Que descrição devo colocar neste campo?"

    # Fallback final sem inglês solto no início
    return f"Que informação queres preencher neste campo: “{raw}”?"

